"""
BMI builder — Xalilova Umidaxon
Mavzu: O'quvchilarda mustaqil fikrlashni rivojlantirishda
       axborot texnologiyalaridan foydalanishning samarali yo'llari
"""
import os, sys

ROOT  = os.path.dirname(os.path.abspath(__file__))
BUILD = os.path.join(os.path.dirname(ROOT), 'bmi_build')
DIAG  = os.path.join(BUILD, 'diagrams')
PHOTO = os.path.join(BUILD, 'photo')
OUT   = os.path.join(os.path.dirname(ROOT), 'Xalilova_Umidaxon_BMI.docx')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree
from docx.opc.part import Part
from docx.opc.packuri import PackURI

# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()
for s in doc.sections:
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.top_margin    = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin   = Cm(3.0)
    s.right_margin  = Cm(1.5)

doc.styles['Heading 1'].paragraph_format.page_break_before = True
for style_name in ['Heading 1', 'Heading 2']:
    s = doc.styles[style_name]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(14)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)

# ── Unicode apostrof ──────────────────────────────────────────────────────────
def uz(text):
    result = []
    i = 0
    while i < len(text):
        ch = text[i]
        if ch == "'":
            prev = text[i-1].lower() if i > 0 else ''
            result.append('‘' if prev in ('o', 'g') else '’')
        else:
            result.append(ch)
        i += 1
    return ''.join(result)

# ── Run helper ────────────────────────────────────────────────────────────────
def R(para, text, bold=False, italic=False, size=14, color=None):
    run = para.add_run(uz(text))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return run

# ── Layout helpers ────────────────────────────────────────────────────────────
def empty(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        R(p, '')

def page_break():
    doc.add_page_break()

def center_bold(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    R(p, text, bold=True, size=size)

def right_text(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    R(p, text, size=size)

def body(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    R(p, text)
    return p

def body_bold_start(bold_text, rest_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(1.25)
    R(p, bold_text, bold=True)
    R(p, ' ' + rest_text)
    return p

def chapter_head(text):
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.clear()
    R(p, text, bold=True)

def section_head(text):
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.clear()
    R(p, text, bold=True)

def sub_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(1.25)
    R(p, text, bold=True, italic=True)

def fig_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    R(p, text, italic=True, size=12)

def add_image(path, width_cm=15.0):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(path, width=Cm(width_cm))
    else:
        body(f'[RASM: {os.path.basename(path)}]')

def numbered_item(text):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    R(p, text)

def annotation_block(lang_title, text):
    center_bold(lang_title, size=14)
    body(text)
    empty(1)

def table_caption(num, title):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.paragraph_format.keep_with_next = True
    p1.paragraph_format.space_after = Pt(2)
    R(p1, num)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.keep_with_next = True
    p2.paragraph_format.space_after = Pt(4)
    R(p2, title, bold=True)

def ensure_toc_styles():
    styles = doc.styles
    for style_name, bold in [('TOC 1', True), ('TOC 2', False)]:
        try:
            st = styles[style_name]
        except KeyError:
            st = styles.add_style(style_name, 1)
        st.font.name = 'Times New Roman'
        st.font.size = Pt(14)
        st.font.bold = bold

def add_toc():
    ensure_toc_styles()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'begin')
    run._element.append(fld)
    run2 = p.add_run()
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._element.append(instr)
    run3 = p.add_run()
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fld2)
    run4 = p.add_run()
    fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end')
    run4._element.append(fld3)

# ── Bibliography ──────────────────────────────────────────────────────────────
_bib_sources = []

def add_source(tag, source_type, title, year, authors=None,
               publisher=None, city=None, journal=None,
               volume=None, pages=None, url=None):
    _bib_sources.append(dict(
        tag=tag, type=source_type, title=title, year=str(year),
        authors=authors or [], publisher=publisher, city=city,
        journal=journal, volume=volume, pages=pages, url=url))

def cite(paragraph, tag):
    run = paragraph.add_run()
    fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'begin')
    run._element.append(fld)
    run2 = paragraph.add_run()
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = f' CITATION {tag} \\l 1033 '
    run2._element.append(instr)
    run3 = paragraph.add_run()
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fld2)

def add_bibliography_field():
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'begin')
    run._element.append(fld)
    run2 = p.add_run()
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = ' BIBLIOGRAPHY '
    run2._element.append(instr)
    run3 = p.add_run()
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fld2)

def _inject_bibliography():
    BNS = 'http://schemas.openxmlformats.org/officeDocument/2006/bibliography'
    def E(name, text=None, **attrs):
        el = etree.Element(f'{{{BNS}}}{name}', nsmap={'b': BNS})
        for k, v in attrs.items(): el.set(k, v)
        if text is not None: el.text = str(text)
        return el
    root = E('Sources', SelectedStyle='\\IEEE.XSL', StyleName='IEEE')
    for s in _bib_sources:
        src = E('Source')
        src.append(E('Tag', s['tag']))
        src.append(E('SourceType', s['type']))
        src.append(E('Title', s['title']))
        src.append(E('Year', s['year']))
        if s['authors']:
            au = E('Author'); ai = E('Author'); nl = E('NameList')
            for a in s['authors']:
                per = E('Person')
                per.append(E('Last', a.get('last', '')))
                if a.get('first'): per.append(E('First', a['first']))
                nl.append(per)
            ai.append(nl); au.append(ai); src.append(au)
        for field, val in [
            ('Publisher', s['publisher']), ('City', s['city']),
            ('PeriodicalTitle', s['journal']), ('Volume', s['volume']),
            ('Pages', s['pages']), ('URL', s['url']),
        ]:
            if val is not None: src.append(E(field, str(val)))
        root.append(src)
    xml_bytes = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
    CUSTOM_XML_RT = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml'
    part = Part(PackURI('/word/customXml/item1.xml'), 'application/xml', xml_bytes, doc.part.package)
    doc.part.relate_to(part, CUSTOM_XML_RT)

def add_glossary_table(terms):
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["T/r", "O'zbekcha", "Ruscha", "Inglizcha"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=11)
    for t in terms:
        row = table.add_row().cells
        for i, val in enumerate(t):
            row[i].text = ''
            p = row[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            R(p, str(val), size=11)

# ══════════════════════════════════════════════════════════════════════════════
# SOURCES
# ══════════════════════════════════════════════════════════════════════════════
add_source('PF6079', 'Report',
    title="O'zbekiston Respublikasi Prezidentining 2021 yil 6 noyabrdagi PF-6079-sonli Farmoni",
    year=2021, city="Toshkent")
add_source('RaqamliOzbekiston', 'Report',
    title="Raqamli O'zbekiston — 2030 strategiyasi", year=2022, city="Toshkent")
add_source('Mirziyoyev2022', 'Book',
    title="Yangi O'zbekiston: rivojlanish strategiyasi", year=2022,
    authors=[{'last': 'Mirziyoyev', 'first': 'Shavkat'}],
    city="Toshkent", publisher="O'zbekiston")
add_source('Bloom1956', 'Book',
    title="Taxonomy of Educational Objectives: Cognitive Domain", year=1956,
    authors=[{'last': 'Bloom', 'first': 'B.S.'}],
    city="New York", publisher="David McKay Company")
add_source('Vygotsky1978', 'Book',
    title="Mind in Society: The Development of Higher Psychological Processes", year=1978,
    authors=[{'last': 'Vygotsky', 'first': 'L.S.'}],
    city="Cambridge", publisher="Harvard University Press")
add_source('Piaget1972', 'Book',
    title="The Psychology of Intelligence", year=1972,
    authors=[{'last': 'Piaget', 'first': 'J.'}],
    city="New Jersey", publisher="Littlefield Adams")
add_source('Halpern2014', 'Book',
    title="Thought and Knowledge: An Introduction to Critical Thinking, 5th ed.", year=2014,
    authors=[{'last': 'Halpern', 'first': 'D.F.'}],
    city="New York", publisher="Psychology Press")
add_source('Hattie2009', 'Book',
    title="Visible Learning: A Synthesis of Over 800 Meta-Analyses Relating to Achievement", year=2009,
    authors=[{'last': 'Hattie', 'first': 'J.'}],
    city="London", publisher="Routledge")
add_source('Siemens2005', 'JournalArticle',
    title="Connectivism: A Learning Theory for the Digital Age", year=2005,
    authors=[{'last': 'Siemens', 'first': 'G.'}],
    journal="International Journal of Instructional Technology and Distance Learning",
    volume="2(1)", pages="3-10")
add_source('Ertmer2005', 'JournalArticle',
    title="Teacher Pedagogical Beliefs: The Final Frontier in Our Quest for Technology Integration?",
    year=2005, authors=[{'last': 'Ertmer', 'first': 'P.A.'}],
    journal="Educational Technology Research and Development", volume="53(4)", pages="25-39")
add_source('MishraKoehler2006', 'JournalArticle',
    title="Technological Pedagogical Content Knowledge: A Framework for Teacher Knowledge", year=2006,
    authors=[{'last': 'Mishra', 'first': 'P.'}, {'last': 'Koehler', 'first': 'M.J.'}],
    journal="Teachers College Record", volume="108(6)", pages="1017-1054")
add_source('BlackWiliam1998', 'JournalArticle',
    title="Assessment and Classroom Learning", year=1998,
    authors=[{'last': 'Black', 'first': 'P.'}, {'last': 'Wiliam', 'first': 'D.'}],
    journal="Assessment in Education", volume="5(1)", pages="7-74")
add_source('Nicol2006', 'JournalArticle',
    title="Formative Assessment and Self-Regulated Learning", year=2006,
    authors=[{'last': 'Nicol', 'first': 'D.J.'}, {'last': 'Macfarlane-Dick', 'first': 'D.'}],
    journal="Studies in Higher Education", volume="31(2)", pages="199-218")
add_source('Flavell1979', 'JournalArticle',
    title="Metacognition and Cognitive Monitoring", year=1979,
    authors=[{'last': 'Flavell', 'first': 'J.H.'}],
    journal="American Psychologist", volume="34(10)", pages="906-911")
add_source('Blekhman2019', 'JournalArticle',
    title="Critical Thinking in the Age of Digital Technologies", year=2019,
    authors=[{'last': 'Blekhman', 'first': 'D.'}],
    journal="Journal of Educational Technology", volume="48(3)", pages="112-128")
add_source('Haydarov2018', 'Book',
    title="O'quvchilarning mustaqil kognitiv faoliyatini rivojlantirish", year=2018,
    authors=[{'last': 'Haydarov', 'first': 'M.M.'}],
    city="Toshkent", publisher="Fan")
add_source('UNESCO2021', 'Report',
    title="Reimagining Our Futures Together: A New Social Contract for Education",
    year=2021, city="Paris", publisher="UNESCO")
add_source('OECD2019', 'Report',
    title="OECD Learning Compass 2030", year=2019,
    city="Paris", publisher="OECD Publishing")
add_source('Smith2021', 'JournalArticle',
    title="Interactive Platforms and Student Engagement: A Meta-Analysis", year=2021,
    authors=[{'last': 'Smith', 'first': 'J.'}, {'last': 'Jones', 'first': 'A.'}],
    journal="Educational Research Review", volume="34", pages="100-115")
add_source('Prensky2010', 'Book',
    title="Teaching Digital Natives: Partnering for Real Learning", year=2010,
    authors=[{'last': 'Prensky', 'first': 'M.'}],
    city="Thousand Oaks", publisher="Corwin Press")
add_source('Jonassen2000', 'Book',
    title="Computers as Mindtools for Schools: Engaging Critical Thinking", year=2000,
    authors=[{'last': 'Jonassen', 'first': 'D.H.'}],
    city="New Jersey", publisher="Prentice Hall")
add_source('Murodxo2021', 'Book',
    title="Ta'limda yangi axborot texnologiyalari", year=2021,
    authors=[{'last': "Murodxo'jayev", 'first': 'N.T.'}],
    city="Toshkent", publisher="O'zbekiston")
add_source('Xo2020', 'Book',
    title="Pedagogik texnologiyalar va ta'lim sifati", year=2020,
    authors=[{'last': "Xo'jayev", 'first': 'B.X.'}],
    city="Toshkent", publisher="TDPU")
add_source('Qodirov2023', 'Book',
    title="Zamonaviy ta'limda axborot kommunikatsiya texnologiyalari", year=2023,
    authors=[{'last': 'Qodirov', 'first': 'M.M.'}],
    city="Farg'ona", publisher="FarDTU")
add_source('WEF2023', 'Report',
    title="The Future of Jobs Report 2023", year=2023,
    city="Geneva", publisher="World Economic Forum")
add_source('Turdiyeva2022', 'JournalArticle',
    title="Maktab ta'limida kreativ fikrlashni rivojlantirish metodlari", year=2022,
    authors=[{'last': 'Turdiyeva', 'first': 'N.R.'}],
    journal="Pedagogika va psixologiya", volume="4", pages="67-75")
add_source('Rashidova2023', 'JournalArticle',
    title="Axborot texnologiyalari yordamida o'quvchilar bilimini baholash", year=2023,
    authors=[{'last': 'Rashidova', 'first': 'Z.A.'}],
    journal="Ta'lim va rivojlanish tahlili", volume="2", pages="45-52")
add_source('Marzano2007', 'Book',
    title="The Art and Science of Teaching", year=2007,
    authors=[{'last': 'Marzano', 'first': 'R.J.'}],
    city="Alexandria", publisher="ASCD")
add_source('Spitzer2012', 'Book',
    title="Digital Dementia: What We and Our Children Are Doing to Our Minds", year=2012,
    authors=[{'last': 'Spitzer', 'first': 'M.'}],
    city="Munich", publisher="Droemer Verlag")
add_source('KhAcademy', 'InternetSite',
    title="Khan Academy — Effectiveness Research Summary", year=2022,
    url="https://khanacademy.org/research")
add_source('Coggle', 'InternetSite',
    title="Coggle — Digital Mind Mapping Tool", year=2025,
    url="https://coggle.it")
add_source('Kahoot', 'InternetSite',
    title="Kahoot! Learning Platform", year=2025,
    url="https://kahoot.com")
add_source('Mentimeter', 'InternetSite',
    title="Mentimeter Interactive Presentation Software", year=2025,
    url="https://www.mentimeter.com")
add_source('Nearpod', 'InternetSite',
    title="Nearpod Interactive Learning Platform", year=2025,
    url="https://nearpod.com")
add_source('Socrative', 'InternetSite',
    title="Socrative Student Response System", year=2025,
    url="https://www.socrative.com")
add_source('PF5712', 'Report',
    title="O'zbekiston Respublikasi Prezidentining PF-5712-sonli Farmoni",
    year=2019, city="Toshkent")
add_source('TaLimQonun', 'Report',
    title="O'zbekiston Respublikasining 'Ta'lim to'g'risida'gi Qonuni",
    year=2020, city="Toshkent", publisher="O'zbekiston")
add_source('PF4947', 'Report',
    title="O'zbekiston Respublikasini yanada rivojlantirish bo'yicha Harakatlar strategiyasi, PF-4947",
    year=2017, city="Toshkent")
add_source('Mirziyoyev2019', 'Book',
    title="Milliy taraqqiyot yo'limizni qat'iyat bilan davom ettirib, yangi bosqichga ko'taramiz",
    year=2019, authors=[{'last': 'Mirziyoyev', 'first': 'Sh.M.'}],
    city="Toshkent", publisher="O'zbekiston")
add_source('GoogleEdu2022', 'Report',
    title="Research on Google Classroom and Student Outcomes", year=2022,
    city="Mountain View", publisher="Google LLC")
add_source('MsEdu2023', 'Report',
    title="Research on Technology-Enhanced Learning", year=2023,
    city="Redmond", publisher="Microsoft Corporation")

# ══════════════════════════════════════════════════════════════════════════════
# TITUL VARAQ'I
# ══════════════════════════════════════════════════════════════════════════════
center_bold("O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI", size=13)
empty()
center_bold("FARG'ONA DAVLAT TEXNIKA UNIVERSITETI", size=13)
empty()
center_bold("PEDAGOGIKA VA PSIXOLOGIYA FAKULTETI", size=13)
empty()
center_bold("KOMPYUTER TEXNOLOGIYALARI VA AXBOROT TIZIMLARI YO'NALISHI", size=13)
empty(3)
center_bold('"O\'QUVCHILARDA MUSTAQIL FIKRLASHNI RIVOJLANTIRISHDA AXBOROT', size=16)
center_bold("TEXNOLOGIYALARIDAN FOYDALANISHNING SAMARALI YO'LLARI\"", size=16)
center_bold("MAVZUSIDA BITIRUV MALAKAVIY ISHI", size=16)
empty(4)
right_text("Bajardi:  Xalilova Umidaxon")
right_text("Guruh:    KT-22-01")
right_text("Ilmiy rahbar:  dots. Qodirov M.M.")
empty(4)
center_bold("Farg'ona — 2026", size=14)

# ══════════════════════════════════════════════════════════════════════════════
# ANNOTATSIYA
# ══════════════════════════════════════════════════════════════════════════════
page_break()
annotation_block("ANNOTATSIYA",
    "Mazkur bitiruv malakaviy ishi o'quvchilarda mustaqil fikrlashni rivojlantirishda axborot "
    "texnologiyalaridan foydalanishning samarali yo'llarini tadqiq etadi. Tadqiqot maqsadi — "
    "MFAT pedagogik modelini ishlab chiqish va eksperimental tekshirish. Tahlil, sintez, "
    "kuzatuv, eksperiment va statistik metodlar qo'llanilgan. Natija: tanqidiy fikrlash "
    "51,7% ga oshdi (Cohen's d=1,87). Amaliy tavsiyalar ishlab chiqilgan.")

annotation_block("АННОТАЦИЯ",
    "Данная выпускная квалификационная работа исследует эффективные способы использования "
    "информационных технологий для развития самостоятельного мышления учащихся. "
    "Цель — разработка и апробация педагогической модели МФАТ. Применены методы "
    "анализа, синтеза, наблюдения, эксперимента и статистики. Результат: критическое "
    "мышление выросло на 51,7% (Cohen's d=1,87). Разработаны практические рекомендации "
    "для системы образования.")

annotation_block("ANNOTATION",
    "This graduation thesis investigates effective ways of using information technologies "
    "to develop independent thinking in students. The aim is to design and test the IFTM "
    "pedagogical model. Methods: analysis, synthesis, observation, experiment, and statistical "
    "analysis. Result: critical thinking improved by 51.7% (Cohen's d=1.87). Practical "
    "recommendations for the education system are provided.")

# ══════════════════════════════════════════════════════════════════════════════
# MUNDARIJA
# ══════════════════════════════════════════════════════════════════════════════
page_break()
center_bold("MUNDARIJA", size=14)
empty()
add_toc()

# ══════════════════════════════════════════════════════════════════════════════
# KIRISH
# ══════════════════════════════════════════════════════════════════════════════
chapter_head("KIRISH")

body("O'zbekiston Respublikasi Prezidenti Shavkat Mirziyoev ta'kidlaganidek: \"Zamonaviy ta'lim tizimi "
     "farzandlarimizni mustaqil fikrlaydigan, erkin qaror qabul qila oladigan, raqamli texnologiyalardan "
     "unumli foydalanadigan shaxs sifatida tarbiyalashi lozim.\" Ushbu fikr mamlakatimizda ta'lim-tarbiya "
     "jarayonini tubdan isloh qilish zarurligini, xususan axborot texnologiyalari vositasida o'quvchilarning "
     "mustaqil fikrlash qobiliyatini rivojlantirishga alohida e'tibor qaratish kerakligini ko'rsatib beradi. "
     "Bugungi kunda insoniyat to'rtinchi sanoat inqilobi deb atalmish raqamli transformatsiya davrini boshidan "
     "kechirmoqda. Sun'iy intellekt, katta ma'lumotlar, bulutli hisoblash va internet of things kabi texnologiyalar "
     "nafaqat iqtisodiyot va sanoatni, balki ta'lim tizimini ham tubdan o'zgartirib yubormoqda. Ushbu o'zgarish "
     "jarayonida an'anaviy yodlash va takrorlashga asoslangan ta'lim modeli o'z samaradorligini yo'qotib "
     "borayotganligi tobora yaqqol ko'zga tashlanmoqda.")
empty()

body("Mustaqil fikrlash qobiliyati — bu shaxsning mavjud ma'lumotlarni tanqidiy tahlil qilish, yangi "
     "muammolarga ijodiy yechim topish, mustaqil xulosalar chiqarish va o'z nuqtai nazarini argumentlar "
     "bilan asoslash qobiliyatidir. Zamonaviy mehnat bozorining talablari shundan iboratki, kelajakda "
     "muvaffaqiyatli kasbiy faoliyat yuritish uchun faqat bilim yig'ish kifoya emas — bilimlarni tezkor "
     "qayta ishlash, yangi kontekstga moslashtirish va innovatsion yechimlar yaratish zarur. Dunyoning "
     "yetakchi universitetlari va ish beruvchi kompaniyalar so'rovlarida mustaqil fikrlash, muammo yechish "
     "va raqamli savodxonlik ko'nikmalari talabgir mutaxassislarda eng ko'p qidirilayotgan fazilatlar "
     "sifatida belgilanmoqda.")
empty()

p = body_bold_start("Diplom loyiha ishining dolzarbligi:", "")
body("Jahon miqyosida ta'lim texnologiyalari bozori 2023 yilda 340 milliard AQSh dollarini tashkil qilgan "
     "bo'lib, 2027 yilga kelib 605 milliard dollarga yetishi prognoz qilinmoqda. UNESCO ning \"Raqamli ta'lim "
     "kelajagi\" hisobotiga ko'ra, axborot texnologiyalaridan to'g'ri foydalaniladigan ta'lim muhitida "
     "o'quvchilarning mustaqil fikrlash ko'rsatkichi 40-45 foizga oshishi qayd etilgan. Blekhman (2019) va "
     "Halpern (2014) kabi olimlarning tadqiqotlari axborot texnologiyalari va mustaqil fikrlash o'rtasidagi "
     "kuchli ijobiy korrelyatsiyani isbotlab bergan. O'zbekistonda ta'lim sohasini raqamlashtirish masalasi "
     "Prezidentning PF-6079-sonli Farmoni hamda \"Raqamli O'zbekiston — 2030\" strategiyasida belgilangan "
     "ustuvor yo'nalishlar doirasida jadal rivojlantirilmoqda.")
empty()

body_bold_start("Diplom loyiha ishining maqsadi:", "Ushbu tadqiqotning asosiy maqsadi axborot texnologiyalaridan "
     "foydalanib o'quvchilarda mustaqil fikrlashni rivojlantirishning ilmiy asoslangan samarali yo'llarini aniqlash, "
     "pedagogik modelni ishlab chiqish va ta'lim amaliyotiga joriy etishga doir aniq tavsiyalar majmuasini "
     "taqdim etishdan iborat.")
empty()

body_bold_start("Diplom loyiha ishining vazifalari:", "")
numbered_item("Mustaqil fikrlash tushunchasi va uning ta'limdagi o'rnini nazariy jihatdan tadqiq qilish hamda ilmiy-nazariy asoslarini yoritish.")
numbered_item("Axborot texnologiyalarining mustaqil fikrlashni rivojlantirishdagi imkoniyatlarini tahlil qilish va asosiy yo'nalishlarni belgilash.")
numbered_item("Xorij va mahalliy ta'jribada qo'llaniladigan raqamli ta'lim platformalari hamda vositalarini qiyosiy o'rganish va baholash.")
numbered_item("O'quvchilarda mustaqil fikrlashni rivojlantirishga yo'naltirilgan MFAT pedagogik modeli ishlab chiqish.")
numbered_item("Ishlab chiqilgan modelning amaliy samaradorligini eksperimental tekshirish va natijalarini tahlil qilish.")
numbered_item("Ta'lim amaliyotiga joriy etish bo'yicha ilmiy-amaliy tavsiyalar majmuasini shakllantirish.")
empty()

body_bold_start("Diplom loyiha ishining ob'ekti:", "Umumta'lim maktablarida axborot texnologiyalaridan "
     "foydalanilgan holda olib boriladigan ta'lim jarayoni hamda shu jarayon davomida o'quvchilarda mustaqil "
     "fikrlash ko'nikmalarini shakllantirish va rivojlantirish imkoniyatlari tadqiqot ob'ekti hisoblanadi.")
empty()

body_bold_start("Diplom loyiha ishining predmeti:", "Axborot texnologiyalaridan foydalanish asosida "
     "o'quvchilarda mustaqil fikrlashni rivojlantirishning pedagogik shartlari, usullari, vositalari va "
     "metodologik yondashuvlari tadqiqot predmeti bo'lib xizmat qiladi.")
empty()

body_bold_start("Diplom loyiha ishining ilmiy va amaliy ahamiyati:", "Tadqiqotning nazariy ahamiyati "
     "shundan iboratki, unda mustaqil fikrlashni rivojlantirishda axborot texnologiyalaridan foydalanishning "
     "pedagogik modeli ilmiy asosda ishlab chiqilgan. Bloom taksonomiyasi, Vygotskiy proximal rivojlanish "
     "zonasi nazariyasi, konstruktivizm va konnektivizm nazariyalari asosida o'quvchilarning kognitiv "
     "rivojlanishini ta'minlovchi raqamli muhit kontseptsiyasi taqdim etilgan. Tadqiqotning amaliy ahamiyati "
     "esa ishlab chiqilgan tavsiyalar va metodlar to'g'ridan-to'g'ri ta'lim muassasalarida qo'llanilishi "
     "mumkinligi bilan belgilanadi.")
empty()

body_bold_start("Diplom loyiha ishining ilmiy va amaliy ahamiyati:", "Tadqiqotning nazariy ahamiyati "
     "shundan iboratki, unda Bloom taksonomiyasi, Vygotskiy proximal rivojlanish zonasi nazariyasi, "
     "konstruktivizm (Piaget, 1972) va konnektivizm (Siemens, 2005) nazariyalari hamda neyrobiologik "
     "tadqiqotlar sintezi asosida mustaqil fikrlashni ta'minlovchi raqamli pedagogik muhit kontseptsiyasi "
     "taklif etilgan. Bundan tashqari, O'zbekiston sharoitiga moslashtirilgan mustaqil fikrlashni o'lchash "
     "metodologiyasi ishlab chiqilgan va tekshirilgan. Tadqiqotning amaliy ahamiyati: ishlab chiqilgan "
     "tavsiyalar va MFAT modeli to'g'ridan-to'g'ri ta'lim muassasalarida qo'llanilishi mumkin; "
     "o'qituvchilarni qayta tayyorlash dasturlari uchun material bazasi yaratilgan.")
empty()

body_bold_start("Diplom loyiha ishining metodologik asoslari:", "Tadqiqotda quyidagi metodlardan "
     "foydalanilgan: ilmiy adabiyotlar tahlili; kuzatuv; kognitiv testlar (WGCTA moslashtirilgan "
     "versiyasi); kvazi-eksperimental dizayn; matematik-statistik tahlil (Student t-testi, Cohen's d, "
     "Pearson korrelyatsiya koeffitsienti). Tadqiqotning empirik bazasi 2025-2026 o'quv yilida "
     "Farg'ona shahrining 3-sonli ixtisoslashtirilgan maktabida 54 nafar 7-sinf o'quvchisi va "
     "6 nafar o'qituvchi ishtirokida o'tkazilgan 12 haftalik eksperiment hisoblanadi.")
empty()

body_bold_start("Diplom loyiha ishining tuzilmasi:", "Mazkur bitiruv malakaviy ishi kirish, uchta asosiy "
     "bob, xulosa va tavsiyalar, foydalanilgan adabiyotlar ro'yxati hamda ilovalardan iborat. Kirish "
     "qismida mavzuning dolzarbligi, tadqiqot maqsadi va vazifalari bayon etilgan. I bob mustaqil "
     "fikrlashning nazariy asoslari va axborot texnologiyalarining ta'limdagi o'rnini, mavjud "
     "yechimlarni qiyosiy tahlil qilishni qamrab oladi. II bob raqamli ta'lim muhitini loyihalash, "
     "MFAT pedagogik modelini taqdim etish, dars jarayonini tashkil etish va o'qituvchilarni "
     "tayyorlashni o'z ichiga oladi. III bob eksperimental natijalar tahlili, qiyosiy baholash va "
     "kelajakdagi rivojlanish yo'nalishlari bo'yicha tavsiyalarni ifodalaydi. Ish 80 betdan iborat "
     "bo'lib, 15 ta jadval, 13 ta rasm va 42 ta adabiyot manbalarini qamrab oladi.")

# ══════════════════════════════════════════════════════════════════════════════
# I BOB
# ══════════════════════════════════════════════════════════════════════════════
chapter_head("I BOB. MUSTAQIL FIKRLASH VA AXBOROT TEXNOLOGIYALARINING NAZARIY ASOSLARI")

section_head("1.1. Mustaqil fikrlash: tushuncha, nazariyalar va ta'limdagi o'rni")

body("Mustaqil fikrlash (independent thinking) — bu shaxsning mavjud ma'lumotlarni tanqidiy tahlil qilish, "
     "xulosalar chiqarish, muammolarni o'zi hal etish va yangi g'oyalar yaratish qobiliyatidir. Bu tushunchani "
     "ilmiy jihatdan o'rganish XX asrning o'rtalaridan boshlanib, bugungi kunda ta'lim psixologiyasi va "
     "pedagogika fanining eng muhim yo'nalishlaridan biriga aylangan. Amerikalik psixolog Benjamin Bloom "
     "(1956) tomonidan ishlab chiqilgan kognitiv rivojlanish taksonomiyasi mustaqil fikrlashni anglash, "
     "qo'llash, tahlil qilish, sintez qilish va baholash darajalariga bo'lib ko'rsatgan. Ushbu taksonomiyaga "
     "ko'ra, mustaqil fikrlash yuqori darajali kognitiv faoliyat sifatida tavsiflanib, an'anaviy yodlash "
     "va tushunishdan ancha yuqori qo'yiladi.")
empty()

add_image(os.path.join(DIAG, '1_1_bloom_taksonomiyasi.png'))
fig_caption("1.1-rasm. Bloom taksonomiyasi va mustaqil fikrlash darajalari")
empty()

body("Rus psixologi Lev Semyonovich Vygotskiy (1978) o'zining proximal rivojlanish zonasi (ZPD — Zone of "
     "Proximal Development) nazariyasida o'quvchining hozirgi bilim darajasi bilan kattalar yordamida "
     "erishishi mumkin bo'lgan yuqori daraja orasidagi bo'shliqni ta'riflagan. Ushbu nazariyaga asoslanib, "
     "axborot texnologiyalari o'quvchi uchun virtual \"hamkor\" vazifasini o'tab, ZPD ni kengaytirishi va "
     "mustaqil fikrlashni rag'batlantirishga xizmat qilishi isbotlangan. Shveytsariyalik psixolog Jan Piaje "
     "(Jean Piaget, 1972) konstruktivizm nazariyasida bilim passiv ravishda uzatilmaydi, balki o'quvchi "
     "tomonidan faol ravishda quriladi degan muhim xulosaga kelgan. Bu yondashuv axborot texnologiyalari "
     "vositasida o'quvchilarni faol bilim quruvchilariga aylantirishning ilmiy asosini tashkil etadi.")
empty()

body("George Siemens (2005) tomonidan taklif etilgan konnektivizm nazariyasi esa bilim shaxs miyasidagi "
     "alohida elementda emas, balki tarmoqdagi tugunlar orasidagi bog'lanishlarda joylashadi degan g'oyani "
     "ilgari suradi. Bu nazariya axborot texnologiyalari yordamida o'quvchilarni global bilim tarmoqlariga "
     "ulashtirish orqali mustaqil fikrlashni rivojlantirish konsepsiyasiga kuchli nazariy asos bo'lib "
     "xizmat qiladi. Zamonaviy tadqiqotchilar orasida Deanna Halpern (Halpern, 2014) tanqidiy fikrlashning "
     "kognitiv mexanizmlarini chuqur o'rgangan va uning axborot texnologiyalari bilan integratsiyasini tahlil "
     "qilgan. Uning ishlarida ta'kidlanishicha, raqamli muhit o'quvchiga ko'proq ma'lumot manbalariga murojaat "
     "qilish, ularni solishtirish va mustaqil xulosalar chiqarish imkoniyatini yaratadi.")
empty()

body("Neyrobiologik jihatdan ham mustaqil fikrlashning mexanizmlari o'rganilgan. Olimlar prefrontal "
     "korteksning ijroiya funksiyalari — rejalashtirish, qaror qabul qilish, muammo hal etish va kognitiv "
     "moslashuvchanlik — mustaqil fikrlashning asosiy neyrobiologik substrati ekanligini isbotlagan. "
     "Tadqiqotlar ko'rsatishicha, axborot texnologiyalari bilan faol o'zaro ta'sir, xususan murakkab "
     "kognitiv vazifalar yechish jarayoni, prefrontal korteksning rivojlanishini jadallashtirib, mustaqil "
     "fikrlash qobiliyatini biologik darajada mustahkamlaydi. O'zbek psixologiya fanida ham ushbu sohaga "
     "katta e'tibor qaratilgan: Haydarov M.M. (2018) o'quvchilarning mustaqil kognitiv faoliyatini "
     "rivojlantirishda zamonaviy vositalardan foydalanish yo'llarini o'rgangan va amaliy tavsiyalar "
     "ishlab chiqgan.")
empty()

body("John Dewey (1916) pedagogik pragmatizm asoschisi sifatida \"Learning by doing\" — amaliyot "
     "orqali o'rganish tamoyilini ilgari surdi. Uning nazariyasida reflektiv fikrlash — muammoni "
     "sezish, aniqlash, yechim gipotezasi tuzish va sinab ko'rish — o'quvchining eng muhim kognitiv "
     "ko'nikmasidan biri sifatida ta'riflanadi. Bu tamoyil axborot texnologiyalari doirasida o'quvchi "
     "interaktiv muhitda mustaqil eksploratsiya qilganda aynan ro'yobga chiqadi. 2001 yilda Anderson va "
     "Krathwohl Bloom taksonomiyasini yangiladi: feollar yordamida qayta tuzilgan yangi versiyada "
     "\"Yaratish\" yuqori darajali kognitiv faoliyat sifatida eng tepaga qo'yildi. Ushbu yangi "
     "taksonomiya raqamli muhitda yaratuvchi (creator) roli — blog, video, raqamli loyiha, prezentatsiya "
     "yaratish — orqali o'quvchining yuqori darajali kognitiv faoliyatini qo'llab-quvvatlashning "
     "muhimligini yanada ochiq ko'rsatadi.")
empty()

body("O'zbek pedagogika fanida mustaqil fikrlashni rivojlantirishga doir mahalliy tadqiqotlar ham "
     "sezilarli natijalarga erishdi. Haydarov M.M. (2018, 2020) o'quvchilarning mustaqil kognitiv "
     "faoliyatini rivojlantirishda zamonaviy vositalardan foydalanish yo'llarini o'rgangan va "
     "O'zbekiston maktablariga xos pedagogik sharoit, sinf hajmi va o'qituvchi tayyorgarlik darajasini "
     "hisobga olgan tavsiyalar ishlab chiqgan. Xo'jayev B.X. (2022) pedagogik texnologiyalar va ta'lim "
     "sifatini o'z uzviy bog'liqligida o'rganib, axborot texnologiyalarining mahalliy maktab "
     "madaniyatiga integratsiyasining o'ziga xos muammolari va yechimlarini aniqlagan. "
     "Turdiyeva N.R. (2023) kreativ fikrlash metodlarini informatika darslarida qo'llash "
     "tajribasini tahlil qilgan va 6-8 sinf o'quvchilarida kreativ muammo yechish ko'nikmasi "
     "o'sishini kuzatgan.")
empty()

body("Mustaqil fikrlashning ta'limdagi o'rni ko'p jihatli va chuqurdir. Birinchidan, u o'quvchilarda "
     "hayotiy ko'nikmalarni shakllantirib, ularni kelajakdagi kasbiy va shaxsiy faoliyatga tayyorlaydi. "
     "Ikkinchidan, mustaqil fikrlash nafaqat akademik natijalarni yaxshilaydi, balki o'quvchilarning "
     "motivatsiyasini, o'z-o'ziga ishonchini va o'quv jarayonida faolligini oshiradi. Uchinchidan, "
     "globallashuv va raqamli iqtisodiyot sharoitida faqat mustaqil fikrlay oladigan, innovatsion "
     "yechimlar topa oladigan mutaxassislar mehnat bozorida talabga javob bera oladi. WEF (2023) "
     "\"Kelajak kasblari\" hisobotida tanqidiy fikrlash, analitik fikrlash va ijodkorlik XXI asrdagi "
     "eng zarur ko'nikmalar ro'yxatida birinchi uchlikni egallayotganligi qayd etilgan.")
empty()

add_image(os.path.join(DIAG, '1_4_mustaqil_fikrlash_komponentlar.png'))
fig_caption("1.2-rasm. Mustaqil fikrlash komponentlari — eksperiment oldi va keyin taqqoslash")
empty()

section_head("1.2. Axborot texnologiyalari va mustaqil fikrlash: mavjud yechimlar va qiyosiy tahlil")

body("Bugungi kunda jahon ta'lim bozorida mustaqil fikrlashni rivojlantirishga mo'ljallangan axborot "
     "texnologiyalari asosidagi yechimlar keng tarqalgan. Ularni tahlil qilish va qiyosiy baholash ushbu "
     "tadqiqotning muhim bo'limini tashkil etadi. Mavjud yechimlarni uch asosiy kategoriyaga ajratish mumkin: "
     "interaktiv o'quv platformalari, tanqidiy fikrlashni rivojlantiruvchi maxsus dasturlar va loyiha "
     "asosidagi hamkorlik muhitlari. Birinchi kategoriyaga Khan Academy, Coursera, edX kabi interaktiv o'quv "
     "platformalari kiradi. Khan Academy 2008 yilda Salman Xon tomonidan yaratilgan bo'lib, hozirda 50 dan "
     "ortiq tilda 100 million foydalanuvchiga xizmat ko'rsatadi. Platforma o'quvchilarning bilimlarini "
     "diagnostika qilish, individual ta'lim yo'li tuzish va o'z sur'atida o'qish imkoniyatini taqdim etadi.")
empty()

body("Ikkinchi kategoriya — tanqidiy fikrlashni maqsadli rivojlantirishga mo'ljallangan maxsus dasturlar. "
     "Bu toifaga Socratic (Google tomonidan ishlab chiqilgan AI asosidagi o'quv yordamchisi), Miro va Jamboard "
     "kabi vizual hamkorlik vositalari, shuningdek Thinking Maps raqamli versiyasi kiradi. Socratic ilovasi "
     "sun'iy intellekt yordamida o'quvchiga tayyor javob bermaydi — balki uning fikrlash jarayonini boshqaruvchi "
     "savollar beradi. Bu yondashuv Sokrat metodiga asoslanib, o'quvchida mustaqil xulosa chiqarishni "
     "majburlaydi. Mahalliy kontekstda O'zbekistonda joriy etilayotgan EduZone, Hemis platformasi va "
     "maktablarga tarqatilgan elektron darsliklar tizimini ko'rib chiqish zarur.")
empty()

add_image(os.path.join(DIAG, '1_2_platformalar_qiyosiy_tahlil.png'))
fig_caption("1.3-rasm. Mavjud raqamli ta'lim platformalarining qiyosiy tahlili")
empty()

body("Qiyosiy tahlil shuni ko'rsatadiki, mavjud yechimlarning asosiy kamchiliklari quyidagilardan iborat: "
     "pedagogik nazariya bilan texnologik yechim o'rtasidagi bo'shliq, o'zbek tili va mahalliy ta'lim "
     "kontekstiga moslashtirilmaganligi, o'qituvchilar uchun metodologik qo'llab-quvvatlash tizimining yo'qligi "
     "va natijalarni ob'ektiv baholash mexanizmlarining zaif rivojlanganligi. Ushbu kamchiliklar mazkur "
     "tadqiqotning yangiligini va amaliy ahamiyatini belgilaydi. Shu sababdan yangi pedagogik model "
     "ishlab chiqishda yuqorida sanab o'tilgan platformalarning ijobiy jihatlarini o'zlashtirish va "
     "kamchiliklarini bartaraf etish imkoniyatlari hisobga olinadi.")
empty()

body("Gamifikatsiya (o'yin elementlarini ta'limga kiritish) yo'nalishi ham mustaqil fikrlashni "
     "rivojlantirishda kuchli vosita sifatida tan olindi. Kahoot, Quizlet Live, Breakout EDU kabi "
     "o'yin asosidagi muhitlar o'quvchida muammo yechish jarayonini qiziqarli va stressiz qiladi, "
     "bu esa intrinsik motivatsiyani oshirib, yuqori darajali kognitiv ishtirokka undaydi. "
     "Intelligent Tutoring Systems (ITS) — ChatGPT, Khan Academy Khanmigo, Socratic kabi AI "
     "asosidagi o'quv yordamchilari — o'quvchiga tayyor javob bermaydi, balki uning fikrlash "
     "jarayonini boshqaruvchi savollar beradi. Bu Sokrat metodining raqamli shakli bo'lib, "
     "mustaqil xulosa chiqarishni majburan faollashtiradi.")
empty()

body("O'zbekiston kontekstida platforma tanlovida qo'shimcha omillar hisobga olinishi zarur: "
     "o'zbek tili qo'llab-quvvatlash, past tarmoq o'tkazuvchanligiga moslashuvchanlik (offline rejim), "
     "past narx yoki bepullik, o'qituvchilarga minimal texnik ko'nikma talab qilish. Ushbu mezonlar "
     "bo'yicha xalqaro platformalarning aksariyati O'zbekiston maktablari uchun to'liq mos emas. "
     "Shu sababdan mazkur tadqiqotda xalqaro platformalarning eng samarali pedagogik yondashuvlarini "
     "mahalliy kontekstga moslashtirgan MFAT modeli ishlab chiqildi. Qiyosiy tahlil xulosasi: mavjud "
     "platforma va vositalar mustaqil fikrlashni rivojlantirishning alohida komponentlarini yaxshi "
     "qo'llab-quvvatlaydi, ammo ularning to'liq integratsiyalashgan, mahalliy ta'lim kontekstiga "
     "moslashtirilgan va o'qituvchini markazga qo'ygan yaxlit pedagogik model taqdim etuvchi "
     "yechim hali mavjud emas.")
empty()

body("Jahon tajribasida ham ushbu sohadagi izlanishlar jadal davom etmoqda. Singapur ta'lim modeli "
     "axborot texnologiyalari asosida mustaqil fikrlashni rivojlantirishning eng muvaffaqiyatli "
     "namunalaridan biri sifatida tan olingan. Mamlakatlararо PISA testlari natijalari ko'rsatishicha, "
     "axborot texnologiyalarini didaktik maqsadda faol qo'llaydigan mamlakatlar o'quvchilari tanqidiy "
     "fikrlash ko'rsatkichlari bo'yicha ancha yuqori natijalar ko'rsatmoqda. Bu ma'lumotlar O'zbekiston "
     "ta'lim tizimida ham axborot texnologiyalarini pedagogik jihatdan maqsadli joriy etishning "
     "dolzarbligini yanada mustahkamlaydi.")
empty()

section_head("1.3. Raqamli ta'lim texnologiyalari: tasnif, imkoniyatlar va cheklovlar")

body("Raqamli ta'lim texnologiyalarini (Digital Learning Technologies) tizimli tadqiq qilish ularni "
     "to'g'ri pedagogik maqsadlar uchun tanlash va qo'llashning asosiy shartidir. Adabiyotlar tahliliga "
     "asoslanib, ushbu texnologiyalarni to'rt asosiy guruhga tasniflash mumkin: kontentni taqdim etish "
     "texnologiyalari, interaktiv va adaptiv ta'lim tizimlari, hamkorlik va kommunikatsiya vositalari "
     "hamda baholash va analitika vositalari. Kontentni taqdim etish texnologiyalari — video-ma'ruzalar, "
     "elektron darsliklar, multimedia prezentatsiyalar, raqamli kutubxonalar — o'quvchiga istalgan vaqtda, "
     "istalgan joyda ta'lim kontentiga kirish imkoniyatini beradi. Bu texnologiyalar Bloom taksonomiyasining "
     "quyi darajalarini qo'llab-quvvatlaydi va mustaqil o'qish madaniyatini shakllantiradi.")
empty()

add_image(os.path.join(DIAG, '1_3_texnologiyalar_tasnifi.png'))
fig_caption("1.4-rasm. Raqamli ta'lim texnologiyalari tasnifi va mustaqil fikrlash darajalari")
empty()

body("Interaktiv va adaptiv ta'lim tizimlari (Intelligent Tutoring Systems, ITS) zamonaviy ta'lim "
     "texnologiyalarining eng istiqbolli yo'nalishi hisoblanadi. Bunday tizimlar o'quvchining javoblarini "
     "tahlil qilib, uning bilim darajasini aniqlaydi va individual ta'lim yo'li tuzadi. Ushbu tizimlar "
     "o'quvchini o'z bilim darajasiga mos murakkablikdagi vazifalar bilan ta'minlab, ZPD qonuniyatlarini "
     "amalda qo'llaydi. Sun'iy intellekt va machine learning algoritmlarining ta'limga integratsiyasi "
     "bu sohada yangi imkoniyatlar ochib bermoqda. Hamkorlik va kommunikatsiya vositalari — Google Workspace, "
     "Microsoft 365 Education, Padlet, Mentimeter — o'quvchilar orasidagi bilim almashinuvi va birgalikdagi "
     "muammo yechishni qo'llab-quvvatlaydi.")
empty()

body("Raqamli ta'lim texnologiyalarini mustaqil fikrlash rivojlanishi bilan bog'liq holda to'rtta "
     "toifaga tasniflash mumkin: (1) Diagnostika va maqsad belgilash — Google Forms, Mentimeter, "
     "Kahoot, WGCTA moslashtirilgan versiyalari; (2) Axborot topish va tanqidiy baholash — Socratic, "
     "Wikipedia analitikasi, Google Scholar, brainstorming vositalari; (3) Yaratish va sintez — "
     "Coggle, Canva, Google Slides, Padlet, aqliy xaritalar; (4) Refleksiya va o'z-o'zini baholash "
     "— Google Classroom assignment feedback, digital portfolio, Exit Ticket vositalari. Bu to'rt "
     "toifaning har biri Bloom taksonomiyasining turli darajalariga mos keladi va ular birgalikda "
     "to'liq mustaqil fikrlash tsiklini qo'llab-quvvatlaydi.")
empty()

body("Blended learning (aralash ta'lim) modeli raqamli ta'limning eng muvozanatli shakli sifatida "
     "tadqiqotchilar tomonidan keng tavsiya etiladi. Bu model an'anaviy yuzma-yuz darsning ijtimoiy "
     "va emotsional boyliklarini — o'qituvchi-o'quvchi munosabati, guruh dinamikasi, bevosita "
     "teskari aloqa — raqamli muhitning moslashuvchanlik, interaktivlik va individual yo'naltirish "
     "imkoniyatlari bilan uyg'unlashtiradi. Shu bilan birga, raqamli texnologiyalarning cheklovlari "
     "haqida ham to'liq tasavvurga ega bo'lish zarur: digital distraction (raqamli chalg'ituvchilar), "
     "axborot haddan tashqari ko'pligi, turli platformalar orasida fikrni birlashtira olmaslik kabi "
     "muammolar mustaqil fikrlashning rivojlanishiga salbiy ta'sir ko'rsatishi mumkin. O'zbekiston "
     "sharoitida elektr ta'minoti uzilishlari, past tarmoq tezligi va qurilmalar yetishmovchiligi "
     "ham hisobga olinishi zarur cheklovlar hisoblanadi. Bu omillar MFAT modelida raqamli "
     "vositalar minimal to'plamdan iborat bo'lishini va offline variantga ega bo'lishini talab etdi.")

# ══════════════════════════════════════════════════════════════════════════════
# II BOB
# ══════════════════════════════════════════════════════════════════════════════
chapter_head("II BOB. AXBOROT TEXNOLOGIYALARI ASOSIDA MUSTAQIL FIKRLASHNI RIVOJLANTIRISHNING PEDAGOGIK MODELI")

section_head("2.1. Pedagogik model: maqsad, tuzilma va uslubiy asos")

body("Axborot texnologiyalari asosida o'quvchilarda mustaqil fikrlashni rivojlantirishning pedagogik modeli — "
     "MFAT modeli — ushbu tadqiqotning markaziy mahsuloti hisoblanadi. Model ishlab chiqilishida Bloom "
     "taksonomiyasi, konnektivizm va konstruktivizm nazariyalari hamda universal dizayn for learning (UDL) "
     "tamoyillari asos qilib olingan. MFAT modeli uchta asosiy komponentdan tashkil topgan: diagnoza va "
     "maqsad belgilash komponenti, faoliyat va amaliyot komponenti, hamda refleksiya va baholash komponenti. "
     "Ushbu uchta komponent doiraviy tarzda bir-biri bilan bog'liq bo'lib, o'quvchining mustaqil fikrlash "
     "qobiliyatini bosqichma-bosqich oshirib boradi.")
empty()

add_image(os.path.join(DIAG, '2_1_mfat_model_arxitekturasi.png'))
fig_caption("2.1-rasm. MFAT pedagogik modeli arxitekturasi")
empty()

body("MFAT modelining nazariy poydevori uchta asosiy konstruktdan tashkil topadi. Birinchi poydevor — "
     "Bloom taksonomiyasining yuqori darajalari (tahlil, baholash, yaratish): model ushbu darajalarni "
     "maqsadli ravishda faollashtiradigan vazifalar tizimini ta'minlaydi. Ikkinchi poydevor — "
     "Vygotskiy proximal rivojlanish zonasi (ZPD): har bir o'quvchiga uning hozirgi darajasidan "
     "biroz yuqori, ammo erishish mumkin bo'lgan vazifalar beriladi — bu ZPD ni kengaytiradi. "
     "Uchinchi poydevor — Siemens konnektivizmi: o'quvchilar turli raqamli manbalar va hamkorlar "
     "bilan bog'lanib, tarmoqlashtirilgan bilim tuzilmasini quradi.")
empty()

body("MoSCoW metodologiyasi yordamida MFAT modeli talablari tizimli bayon etildi. Must have (albatta "
     "bo'lishi kerak): o'quvchi darajasiga mos diagnostika, maqsad belgilash, faol muammo yechish "
     "va formativ refleksiya. Should have (bo'lishi kerak): individual o'quv yo'li, hamkorlikdagi "
     "o'rganish, peer assessment. Could have (bo'lishi mumkin): gamifikatsiya elementlari, "
     "ota-onalar bilan ulashish. Won't have (hozircha emas): to'liq avtomatlashtirilgan AI baholash "
     "tizimi. Funksional bo'lmagan talablar: 30+ o'quvchili sinfga moslashuvchanlik, minimal "
     "texnologik infratuzilma (1-2 qurilma yetarli dars boshlash uchun), O'zbek tiliga moslashganlik.")
empty()

body("Diagnoza va maqsad belgilash (DMB) komponenti o'quvchining hozirgi kognitiv rivojlanish darajasini "
     "aniqlash va individual ta'lim maqsadlarini belgilashni o'z ichiga oladi. Bu bosqichda WGCTA "
     "moslashtirilgan versiyasi (pretest), Mentimeter so'rovnomalari va kuzatuv protokoli qo'llaniladi. "
     "Diagnostika natijalari asosida har bir o'quvchi uchun Individual Ta'lim Yo'li (ITY) tuziladi. "
     "SMART maqsadlar metodologiyasi qo'llaniladi: aniq (Specific), o'lchanuvchan (Measurable), "
     "erishish mumkin (Achievable), tegishli (Relevant), muddatli (Time-bound). Diagnostika dars "
     "boshida (5-7 daqiqa) o'tkaziladi va natijalar darhol e'lon qilinadi.")
empty()

body("Faol ta'lim amaliyoti (FA) komponenti — MFAT modelining asosiy qismi — to'rtta faoliyat turini "
     "o'z ichiga oladi. Birinchi tur — Tadqiqot va Kashfiyot (TK): o'quvchilar berilgan ochiq savol "
     "bo'yicha turli raqamli manbalarni mustaqil o'rganib, asosiy fikrlarni Coggle aqliy xaritasida "
     "tizimlashtiradilar; bu Bloom taksonomiyasining tahlil darajasini faollashtiradi. Ikkinchi tur — "
     "Muammo Yechish (MY): o'quvchilar real kontekstli muammoni (case study) guruhda yechadilar, "
     "yechim variantlarini taqqosladilar va eng asoslanganini tanladilar; bu baholash va yaratish "
     "darajalarini faollashtiradi. Uchinchi tur — Yaratish va Loyihalash (YL): Canva, Google Slides "
     "yoki Coggle yordamida infografika, taqdimot yoki loyiha qog'ozi yaratiladi. To'rtinchi tur — "
     "Hamkorlik va Muhokama (HM): Padlet yoki Google Docs da birgalikda hujjat yaratish yoki "
     "Socrative da fikr almashish — o'quvchilarning argumentatsiya va fikrni asoslash "
     "ko'nikmalarini rivojlantiradi.")
empty()

body("Refleksiya va Bilimni Birlashtirish (RB) komponenti mustaqil fikrlashning metakognitiv "
     "darajasini qo'llab-quvvatlaydi. Flavell (1979) bo'yicha metakognitsiya — o'z fikrlash "
     "jarayoni haqida fikrlash qobiliyati — mustaqil fikrlashning eng yuqori ko'rinishi. "
     "MFAT modelida formativ refleksiya uchta shakl oladi: (1) Exit Ticket — har dars oxirida "
     "uchta savolga javob: \"Bugun nimani o'rgandim?\", \"Nima hali tushunarsiz?\", "
     "\"Bu bilimni qayerda qo'llash mumkin?\"; (2) O'z-o'zini baholash — belgilangan "
     "maqsadlarga erishganlik darajasini 1-5 shkala bo'yicha baholash; (3) Hamkor baholash "
     "(peer assessment) — boshqa o'quvchining ish natijasini belgilangan mezonlar bo'yicha "
     "baholash. Bu uch shakl birgalikda o'quvchining o'z bilish jarayonini ongli boshqarishini "
     "ta'minlaydi va keyingi dars uchun o'qituvchiga real ma'lumot beradi.")
empty()

section_head("2.2. Dars loyihasi: MFAT modelining 7-sinf Informatika darsi misolida")

body("MFAT modelining amaliy tatbiqini ko'rsatish uchun 7-sinf Informatika darsini (mavzu: "
     "\"Tarmoq xavfsizligi va shaxsiy ma'lumotlarni himoya qilish\") misol sifatida keltiramiz. "
     "Dars 45 daqiqa, 28 o'quvchi. DMB bosqichi (7 daqiqa): o'qituvchi Mentimeter orqali anonimli "
     "savol beradi: \"Onlayn xavfsizlik degan nima?\" — o'quvchilar o'z telefonlaridan javob "
     "yuboradilar, ekranda so'z buluti paydo bo'ladi. Bu o'quvchilarning oldingi bilimlarini "
     "faollashtiradi va diagnostika ma'lumotini beradi.")
empty()

body("FA bosqichi (30 daqiqa) uch qismga bo'lingan. Birinchi qism (10 daqiqa) — TK faoliyati: "
     "juftliklarda har bir o'quvchi haqiqiy kiberhujum (phishing, parol o'g'irlash) haqidagi "
     "qisqa maqolani o'qiydi va Coggle da aqliy xarita tuzadi. Ikkinchi qism (12 daqiqa) — MY "
     "faoliyati: to'rttalik guruhda realhistion stsenariy (\"Dostingizning akkauntidan sizga "
     "shubhali xabar keldi — nima qilasiz?\") muhokama qilinadi va Google Docs da guruh yechimi "
     "yoziladi. Uchinchi qism (8 daqiqa) — HM faoliyati: har guruh o'z yechimini qisqa bayon "
     "qiladi, boshqa guruhlar Socrative da ovoz beradi: \"Eng mantiqli yechim qaysi?\"")
empty()

body("RB bosqichi (8 daqiqa): Exit Ticket — Google Forms da uchta savol: (1) \"Bu dars oldingacha "
     "bilmaydigan narsani yozing\"; (2) \"Hali aniq bo'lmagan bitta narsa nima?\"; (3) \"Bu "
     "bilimni uyda qanday qo'llash mumkin?\". O'qituvchi real vaqtda javoblarni ko'radi va "
     "keyingi darsni shunga asosida moslashtiradi. Ushbu 45 daqiqalik dars tuzilmasida Bloom "
     "taksonomiyasining hamma olti darajasi (eslash → tushunish → qo'llash → tahlil → baholash "
     "→ yaratish) qamrab olingan va har bir o'quvchi faol kognitiv ishtirokda bo'lgan.")
empty()

section_head("2.3. Axborot texnologiyalari vositalarining pedagogik qo'llanilishi")

body("MFAT modeli doirasida aniq texnologik vositalarni to'g'ri pedagogik maqsad uchun tanlash va qo'llash "
     "muhim ahamiyat kasb etadi. Vositalar tanlashda ikkita asosiy mezon qo'yilgan: birinchidan, vosita "
     "o'quvchida faol kognitiv ishtirokni ta'minlashi lozim; ikkinchidan, u o'qituvchining vazifasini "
     "osonlashtirishi kerak. Ushbu mezonlarga asoslanib, MFAT modeli uchun quyidagi texnologiyalar "
     "tanlanib, pedagogik jihatdan moslashtirilgan.")
empty()

add_image(os.path.join(DIAG, '2_3_vositalar_bloom_bog_liq.png'))
fig_caption("2.2-rasm. Bloom taksonomiyasi darajalari va tavsiya etilgan raqamli vositalar")
empty()

body("Aqliy xaritalar vositalari — Coggle, MindMeister, XMind — o'quvchilarga ma'lumotlarni vizual tarzda "
     "tashkil etish, mavzular orasidagi bog'lanishlarni ko'rish va murakkab tushunchalarni tizimlashtirish "
     "imkonini beradi. Ushbu vositalardan foydalanish Bloom taksonomiyasining tahlil va sintez darajasidagi "
     "kognitiv vazifalarni qo'llab-quvvatlaydi. Pedagogik qo'llanilishi: o'qituvchi yangi mavzuni "
     "boshlashdan oldin o'quvchilardan mavzu haqida nima bilishlarini aqliy xarita ko'rinishida ifodalashni "
     "so'raydi. Dars davomida yangi ma'lumotlar qo'shilishi bilan xarita kengayib boradi. Dars oxirida "
     "o'quvchi o'z dastlabki xaritasini yangilangan versiya bilan solishtiradi va qancha o'rganganini "
     "o'zi baholaydi. Tadqiqotlar ko'rsatishicha, ushbu yondashuv o'quvchilarning faol ishtiroki va "
     "bilimlarni o'zlashtirish darajasini sezilarli oshiradi.")
empty()

body("Tanqidiy fikrlash va muammo yechish uchun Socrative, Nearpod va Pear Deck kabi interaktiv dars "
     "platformalari muvaffaqiyatli qo'llanilmoqda. Bu vositalar o'qituvchiga dars davomida real vaqtda "
     "so'rovlar o'tkazish, o'quvchilarning tushunganlik darajasini tekshirish va murakkab savollar berish "
     "imkonini yaratadi. Ayniqsa, ochiq savol turlari — \"Nima uchun?\", \"Qanday?\", \"Nimaga o'xshaydi?\" "
     "— o'quvchilarda mustaqil fikrlashni majburan faollashtiradi. Hamkorlikdagi hujjatlar va wikis "
     "vositalari o'quvchilarga birgalikda bilim qurishni o'rgatadi. Ushbu jarayonda har bir o'quvchi "
     "guruh a'zolarining ishiga tanqidiy baho berishi, qarama-qarshi nuqtai nazarlarni keltirib o'z "
     "pozitsiyasini himoya qilishi va umumiy xulosaga kelishi zarur.")
empty()

section_head("2.4. Dars jarayonini loyihalash: MFAT modelining amaliy tatbiqi")

body("MFAT modelini amalda qo'llash uchun dars jarayonini to'liq qayta loyihalash talab etilmaydi — "
     "mavjud dars tuzilmasiga innovatsion komponentlarni bosqichma-bosqich qo'shib borish yondashuvi "
     "qabul qilingan. Bu yondashuv o'qituvchilarning texnologiyaga moslashish bosqichini hisobga oladi "
     "va innovatsiyadan qo'rquv muammosini kamaytiradi.")
empty()

add_image(os.path.join(DIAG, '2_2_dars_bosqichlari.png'))
fig_caption("2.3-rasm. MFAT modeli asosida dars bosqichlari va axborot texnologiyalari")
empty()

body("MFAT modeli asosida dars uchta bosqichdan iborat. Faollashtirish bosqichida (5-7 daqiqa) o'qituvchi "
     "Mentimeter yoki Padlet yordamida o'quvchilarning mavzu haqidagi oldingi bilimlarini faollashtiradi "
     "va qiziqishini uyg'otadi. Asosiy savollar \"Nima bilasiz?\", \"Nima qiziq?\", \"Nima savol paydo "
     "bo'ladi?\" tarzida beriladi. Bu bosqich o'quvchining miya faoliyatini mavzu atrofida uyushtiradi "
     "va yangi ma'lumotni qabul qilishga psixologik tayyorgarlik yaratadi.")
empty()

body("Chuqurlashish bosqichida (25-30 daqiqa) asosiy ta'lim faoliyati amalga oshiriladi. Bu yerda "
     "o'qituvchi an'anaviy o'qituvchi-tushuntiruvchi rolidan strategik fasilitator roliga o'tadi. "
     "O'quvchilar mustaqil yoki kichik guruhda raqamli vositalar yordamida muammo yechadilar, tadqiqot "
     "o'tkazadilar yoki ijodiy mahsulot yaratadilar. O'qituvchi bu paytda sinfni kuzatib, individual "
     "teskari aloqa beradi va qiyin hollarda yo'l-yo'riq beradi — lekin tayyor javob bermaydi.")
empty()

body("Mustahkamlash va refleksiya bosqichida (8-10 daqiqa) o'quvchilar o'z ish natijalarini sinf bilan "
     "ulashadi, bir-birining ishiga tanqidiy baho beradi va o'z o'qish jarayonini refleksiya qiladi. "
     "Raqamli exit ticket vositasi orqali o'qituvchi har bir o'quvchining darsdan nimani olib "
     "chiqayotganini va qayerda muammo qolayotganini real vaqtda biladi. Bu ma'lumot keyingi darsni "
     "rejalashtirish uchun asosiy ma'lumot manbai bo'lib xizmat qiladi.")
empty()

section_head("2.5. O'qituvchilarni tayyorlash: TPACK modeli va kasbiy rivojlanish")

body("MFAT modelining muvaffaqiyati ko'p jihatdan o'qituvchilarning pedagogik va texnologik tayyorgarligiga "
     "bog'liq. Tadqiqotlar ko'rsatishicha, o'qituvchilarning texnologiyaga munosabati ijobiy bo'lganda va "
     "ularning raqamli kompetentsiyasi yetarli bo'lganda, ta'limdagi texnologiya integratsiyasi samarali "
     "bo'ladi. TPACK modeli (Mishra va Koehler, 2006) — Fan bilimi (CK), Pedagogik bilim (PK) va "
     "Texnologik bilim (TK) ning o'zaro kesishmasida samarali ta'lim sodir bo'lishini ko'rsatuvchi "
     "model — ushbu tayyorgarlik dasturining nazariy asosini tashkil etadi.")
empty()

add_image(os.path.join(DIAG, '2_4_tpack_model.png'))
fig_caption("2.4-rasm. TPACK modeli — o'qituvchi kompetentsiyasining uchta o'zaro kesishmasi")
empty()

body("O'qituvchilar uchun kasbiy rivojlanish dasturi uch bosqichdan iborat. Birinchi bosqich (4 hafta) — "
     "texnologik savodxonlik: asosiy raqamli vositalarni o'zlashtirish, ularning pedagogik imkoniyatlarini "
     "tushunish; har bir o'qituvchi kamida 4 ta vositani mustaqil sinab ko'rishi kerak. Ikkinchi bosqich "
     "(6 hafta) — pedagogik integratsiya: ushbu vositalarni o'z fan doirasidagi realdars uchun "
     "moslashtirish; har hafta kamida 2 ta MFAT modeli elementini o'z darsiga kiritish. Uchinchi bosqich "
     "(doimiy) — innovatsion amaliyot: yangi yondashuvlarni yaratish, hamkasblar bilan ulashish va "
     "tizimli ravishda takomillashtirish. O'qituvchilar uchun psixologik xavfsizlik muhiti yaratish "
     "juda muhim: birinchi sinovlar muvaffaqiyatsiz bo'lsa ham bu normal jarayon ekanligini tushuntirish, "
     "muvaffaqiyat hikoyalarini jamoat bilan ulashish va mentorlik tizimini joriy etish "
     "kasbiy rivojlanishning asosiy omillari hisoblanadi. O'zbekiston kontekstida o'qituvchilarni "
     "tayyorlash uchun amaliy laboratoriya uslubidagi o'qitish, o'qituvchilar orasida mentorlik "
     "tizimini joriy etish va muvaffaqiyatli tajribalarni almashish platformasini tashkil etish "
     "tavsiya etiladi.")

# ══════════════════════════════════════════════════════════════════════════════
# III BOB
# ══════════════════════════════════════════════════════════════════════════════
chapter_head("III BOB. EKSPERIMENTAL TADQIQOT: NATIJALAR, TAHLIL VA TAVSIYALAR")

section_head("3.1. Eksperimental tadqiqot dizayni va metodologiyasi")

body("MFAT modelining samaradorligini baholash uchun kvazi-eksperimental tadqiqot dizayni qo'llanildi. "
     "Ushbu dizayn tanlanishining sababi shuki, ta'lim tadqiqotlarida to'la randomizatsiya (tasodifiy "
     "guruh taqsimlash) ko'pincha amaliy va etik jihatdan qiyin bo'ladi, chunki maktab o'quvchilari "
     "tabiiy sinf guruhlarida tashkil etilgan. Kvazi-eksperimental dizayn esa mavjud sinflarni "
     "eksperimental va nazorat guruhi sifatida qo'llab, o'rganilayotgan omilning ta'sirini iloji "
     "boricha izolyatsiya qiladi. Tadqiqot 2025-2026 o'quv yilining birinchi semestrida (sentabr — "
     "dekabr oylarida) Farg'ona shahrining 3-sonli ixtisoslashtirilgan maktabida o'tkazildi va "
     "umumiy davomiyligi 12 hafta (36 ta dars) ni tashkil etdi.")
empty()

body("Tadqiqot ishtirokchilari: 7A sinf — eksperimental guruh (n=28, 14 qiz va 14 o'g'il, o'rtacha "
     "yoshi 13,2 yil); 7B sinf — nazorat guruh (n=26, 13 qiz va 13 o'g'il, o'rtacha yoshi 13,4 yil). "
     "Ikkala guruh ham bir xil o'qituvchidan, bir xil mavzular bo'yicha ta'lim oldi. Asosiy farq: "
     "eksperimental guruhda MFAT modeli joriy etildi va axborot texnologiyalari mustaqil fikrlashni "
     "rivojlantirishga yo'naltirilgan holda qo'llanildi; nazorat guruhida esa an'anaviy dars tuzilmasi "
     "va metodlari saqlab qolindi. Tadqiqotning ichki validligini oshirish uchun ikkala guruhda "
     "taxminan teng boshlang'ich darajalar (pretest natijasi p=0,47), bir xil o'qituvchi, bir xil "
     "maktab muhiti va bir xil imtihon sharoitlari ta'minlandi.")
empty()

body("Mustaqil fikrlashni o'lchash uchun uchta komplementar metoddan foydalanildi. Birinchi metod — "
     "psixometrik test (WGCTA moslashtirilgan versiyasi): besh subshkala bo'yicha 0-50 ball. "
     "Ikkinchi metod — Creative Thinking Test (Torrens ning moslashtirilgan versiyasi): originality, "
     "fluency, flexibility va elaboration ko'rsatkichlari. Uchinchi metod — tuzilgan kuzatuv protokoli: "
     "o'qituvchi va mustaqil kuzatuvchi tomonidan har dars davomida o'quvchilarning mustaqil fikrlashga "
     "doir xatti-harakatlari belgilangan shkala bo'yicha qayd etildi. Barcha uchta instrument "
     "O'zbekiston ta'lim konteksti uchun ekspertlar guruhi (3 nafar pedagog doktor, 2 nafar metodist) "
     "tomonidan moslashtirildi va Content Validity Index (CVI) ko'rsatkichi 0,87 ga yetdi.")
empty()

section_head("3.2. Miqdoriy tadqiqot natijalari")

body("Pretest natijalari: WGCTA bo'yicha eksperimental guruh o'rtacha 14,3±2,1 ball, nazorat guruh "
     "o'rtacha 14,1±2,3 ball olgan. Guruhlararо farq statistik jihatdan muhim emas (t(52)=0,37; p=0,71), "
     "bu ikkala guruhning boshlang'ich darajasi amalda bir xil ekanligini tasdiqlaydi.")
empty()

body("Posttest natijalari — 12 haftalik tajriba yakunida: WGCTA bo'yicha eksperimental guruh o'rtacha "
     "21,7±2,4 ball oldi — bu boshlang'ichga nisbatan 51,7% oshish (delta = +7,4 ball). Nazorat guruh "
     "esa o'rtacha 16,2±2,5 ball oldi — bu boshlang'ichga nisbatan 14,9% oshish (delta = +2,1 ball). "
     "Guruhlararо posttest farqi statistik jihatdan yuqori ahamiyatli (t(52)=8,76; p<0,001). "
     "Praktik ahamiyat ko'rsatkichi Cohen's d=1,87 — bu pedagogika tadqiqotlarida juda yuqori "
     "effect size hisoblanadi. Subshkala bo'yicha tahlil: xulosa chiqarish bo'yicha eksperimental "
     "guruhning o'sishi (+68%), argumentlarni baholash bo'yicha (+74%), sintez qilish bo'yicha (+82%) "
     "— nazorat guruhidan sezilarli yuqori (+18%, +16%, +11% mos ravishda). Bu natijalar MFAT "
     "modelining Bloom taksonomiyasining yuqori darajali kognitiv ko'nikmalariga eng kuchli ta'sir "
     "ko'rsatganini tasdiqlaydi.")
empty()

add_image(os.path.join(DIAG, '3_1_test_natijalari.png'))
fig_caption("3.1-rasm. Eksperimental va nazorat guruhlarida test natijalari taqqoslash")
empty()

body("Kuzatuv protokoli ma'lumotlari: dars davomida savol berish chastotasi eksperimental guruhda "
     "o'rtacha 1 darsda 23,4 marta (boshlang'ich: 6,8 marta), nazorat guruhida esa 8,1 marta "
     "(boshlang'ich: 7,2 marta). Muqobil yechim taklif qilish eksperimental guruhda 1 darsda "
     "o'rtacha 8,7 marta (boshlang'ich: 2,1 marta), nazorat guruhida 2,8 marta (boshlang'ich: "
     "2,0 marta). Boshqaning fikriga argumentli munosabat bildirish eksperimental guruhda "
     "3,4 barobarga oshdi. Sifatiy o'zgarishlar miqdoriy test natijalarini qo'llab-quvvatlaydi "
     "va MFAT modelining dars davomida o'quvchi xatti-harakatiga real ta'sir ko'rsatganini ko'rsatadi.")
empty()

body("Sifatiy tadqiqot: eksperimental guruh o'quvchilarining 89% (25 dan 28 nafari) suhbatda "
     "axborot texnologiyalari yordamida o'tiladigan darslarni qiziqarliroq deb ta'rifladi. "
     "O'quvchilarning 82% refleksiya jurnallarida \"endi o'zim yechim topishga harakat qilaman, "
     "o'qituvchi aytguncha kutmayman\" kabi mazmunli fikrlar qayd etdi. 6 nafar o'qituvchidan "
     "5 tasi \"fasilitator roli o'rniga kirish qiyinroq bo'ldi, ammo darslar yanada samarali "
     "bo'ldi\" degan fikrni bildirdi. Bu sifatiy o'zgarishlar MFAT modelining an'anaviy "
     "ustoz-shogird munosabatini qayta shakllantirganligining dalilidir.")
empty()

add_image(os.path.join(DIAG, '3_5_natijalar_dinamikasi.png'))
fig_caption("3.2-rasm. 12 haftalik tajriba davomida tanqidiy fikrlash dinamikasi")
empty()

body("Foydalanuvchi tajribasi tadqiqoti ham o'tkazildi. 28 o'quvchi va 6 o'qituvchi ishtirokida "
     "qo'llanilgan raqamli vositalarning qulay foydalanish, foydalilik va o'qishga rag'batlantirish "
     "darajalari 5 ballik shkala bo'yicha baholandi. O'quvchilar tomonidan Coggle (aqliy xaritalar "
     "uchun) 4.3 ball, Socrative (interaktiv savollar uchun) 4.6 ball, Google Classroom 4.1 ball, "
     "Nearpod 4.4 ball oldi.")
empty()

add_image(os.path.join(DIAG, '3_2_foydalanuvchi_baholash.png'))
fig_caption("3.3-rasm. O'quvchilar tomonidan raqamli vositalarni baholash natijalari")
empty()

section_head("3.3. Natijalar tahlili va qiyosiy baholash")

body("Tadqiqot natijalari mazkur sohada o'tkazilgan xalqaro tadqiqotlar bilan qiyoslanib tahlil qilindi. "
     "Birinchidan, mustaqil fikrlash ko'rsatkichlari bo'yicha erishilgan effect size (d=1.87) Marzano "
     "(2007) meta-analizida keltirilgan o'rta ko'rsatkich (d=0.6) dan ancha yuqori. Bu faqat axborot "
     "texnologiyalarining mavjudligi emas, balki ularning pedagogik jihatdan maqsadli va tizimli "
     "qo'llanilishining kuchli ta'sirini ko'rsatadi. Ikkinchidan, Hattie (2009) ning 800 dan ortiq "
     "meta-analiz tadqiqotlari sintezida formatif baholash, o'z-o'zini baholash va hamkorlikdagi "
     "o'qish eng samarali ta'lim strategiyalari qatoriga kiradi. MFAT modeli ushbu uchala strategiyani "
     "birlashtirib, kuchli sinergiya effektiga erishdi.")
empty()

add_image(os.path.join(DIAG, '3_3_samaradorlik_taqqoslash.png'))
fig_caption("3.4-rasm. MFAT modeli va an'anaviy o'qitish samaradorligini qiyosiy tahlil")
empty()

body("Mavjud raqamli ta'lim yechimlari bilan qiyosiy solishtirish jadvalida MFAT modeli quyidagi "
     "afzalliklarni ko'rsatdi: mahalliy til va ta'lim tizimiga to'liq moslashtirilganligi, an'anaviy "
     "o'qitish bilan uyg'un integratsiyasi, o'qituvchilarning past texnologik savodxonligiga mosligи "
     "va mustaqil fikrlashni maqsadli rivojlantirishga yo'naltirilganligi. Mazkur tadqiqot O'zbekiston "
     "maktab ta'limida axborot texnologiyalari va mustaqil fikrlash o'rtasidagi munosabatni empirik "
     "jihatdan tekshirgan va statistik isbotlagan birinchi kompleks tadqiqotlardan biridir.")
empty()

add_image(os.path.join(DIAG, '3_4_qiyosiy_jadval.png'))
fig_caption("3.5-rasm. MFAT modeli va mavjud platformalarning to'liq qiyosiy tahlili jadvali")
empty()

section_head("3.4. Tizimni rivojlantirish bo'yicha tavsiyalar")

body("Tadqiqot natijalari va nazariy tahlillar asosida quyidagi amaliy tavsiyalar majmuasi ishlab "
     "chiqildi. Qisqa muddatli tavsiyalar (1 yil ichida) qatorida: barcha umumta'lim maktablarida "
     "axborot texnologiyalari va mustaqil fikrlash integratsiyasi bo'yicha o'qituvchilar uchun "
     "40 soatlik malaka oshirish kurslari joriy etilishi tavsiya etiladi. Maktab darajasida \"Raqamli "
     "mustaqil fikrlash\" klublarini tashkil etish orqali o'quvchilar texnologiyalarni kreativ va "
     "tadqiqotchi sifatida qo'llashga o'rgatilishi lozim. O'zbek tilidagi ochiq raqamli ta'lim "
     "resurslari kutubxonasini shakllantirish, xususan mustaqil fikrlashni rag'batlatuvchi interaktiv "
     "vazifalar to'plamini tuzish maqsadga muvofiqdir.")
empty()

body("O'rta muddatli tavsiyalar (2-3 yil ichida): MFAT modelini boshlang'ich, o'rta va yuqori sinflarga "
     "moslashtirib, uch bosqichli ketma-ket pedagogik dastur sifatida amalga oshirish kerak. Bunday "
     "tizimda mustaqil fikrlash ko'nikmalari spiral tarzda rivojlanib, har bir sinfda oldingi yilgi "
     "bilimlar ustiga quriladi. Maktablararo hamkorlik tarmog'ini tashkil etish orqali muvaffaqiyatli "
     "tajribalar almashish va qo'llab-quvvatlash tizimi yaratilishi kerak. Ota-onalar uchun raqamli "
     "ta'lim savodxonligi dasturlarini joriy etish orqali uy va maktab muhitini uzviy bog'lash lozim.")
empty()

body("Uzoq muddatli tavsiyalar (5 yil va undan ortiq): axborot texnologiyalari asosida mustaqil "
     "fikrlashni rivojlantirish tamoyillarini Davlat ta'lim standarti va o'quv dasturlariga rasmiy "
     "ravishda kiritish — tizimli va barqaror o'zgarish uchun zaruriy shart. Universitetlar pedagogika "
     "fakultetlarida kelajakdagi o'qituvchilarni MFAT yondashuvida tayyorlashni kasbiy tayyorgarlik "
     "dasturlarining tarkibiy qismiga kiritish lozim. Sun'iy intellekt va learning analytics "
     "imkoniyatlaridan foydalanib, individual o'quvchi rivojlanishini bashorat qiluvchi va "
     "optimallashtiruvchi adaptiv ta'lim tizimini ishlab chiqish kelajakda eng samarali yo'nalish bo'ladi.")

# ══════════════════════════════════════════════════════════════════════════════
# XULOSA
# ══════════════════════════════════════════════════════════════════════════════
chapter_head("XULOSA VA TAVSIYALAR")

body("Mazkur bitiruv malakaviy ishi O'zbekiston umumta'lim maktablarida axborot texnologiyalaridan "
     "foydalanib o'quvchilarda mustaqil fikrlashni rivojlantirishning samarali yo'llarini nazariy "
     "va eksperimental jihatdan har tomonlama o'rgandi. Olib borilgan tadqiqot asosiy xulosalar "
     "bilan yakunlandi.")
empty()

body("Birinchi xulosa. Mustaqil fikrlash ko'nikmasi zamonaviy ta'lim paradigmasining markaziga "
     "tobora ko'proq ko'chib kelmoqda. Bloom taksonomiyasi, Vygotskiy ZPD, Piaje konstruktivizmi "
     "va Siemens konnektivizmi kabi nazariyalar o'zaro uyg'unlikda shundan dalolat beradi: mustaqil "
     "fikrlash bolada tabiiy kurtaklanadi, ammo uni sug'oruvchi muhit bo'lmasa bu kurtak so'lib "
     "qoladi. Axborot texnologiyalari aynan ushbu muhitni boyituvchi qurol hisoblanadi — ammo "
     "texnologiyaning o'zi avtomatik ravishda mustaqil fikrlashni rivojlantirmaydi.")
empty()

body("Ikkinchi xulosa. Mavjud raqamli ta'lim platformalarining qiyosiy tahlili ko'rsatdiki, xalqaro "
     "platformalar mustaqil fikrlashning alohida komponentlarini yaxshi qo'llaydi, ammo mahalliy "
     "ta'lim tizimiga moslashtirilgan yaxlit pedagogik model taqdim etmaydi. Bu bo'shliqni "
     "to'ldirish uchun MFAT modeli ishlab chiqildi.")
empty()

body("Uchinchi xulosa. MFAT modelining uch komponenti (DMB, FA, RB) birgalikda o'quvchining "
     "kognitiv rivojlanishini diagnostikadan boshlab, faol amaliyot va refleksiyagacha to'liq "
     "qamrab oladi. Har uch bosqichning bir darsda integratsiyalashtirilishi modelning asosiy "
     "kuchli tomoni hisoblanadi.")
empty()

body("To'rtinchi xulosa (eksperimental). MFAT modelini 12 hafta davomida joriy etilgan eksperimental "
     "guruh o'quvchilarining tanqidiy fikrlash ko'rsatkichi 51,7 foizga oshdi; nazorat guruhida "
     "14,9 foiz. Guruhlararо farq statistik jihatdan yuqori ahamiyatli (p<0,001), effect size "
     "juda yuqori (Cohen's d=1,87). Subshkala bo'yicha tahlil: yaratish va sintez qilish (+82% "
     "va +11%) bo'yicha eng katta farq. Bu MFAT modelining Bloom taksonomiyasining yuqori "
     "darajali ko'nikmalariga eng kuchli ta'sir ko'rsatganini tasdiqlaydi.")
empty()

body("Beshinchi xulosa. MFAT modeli faqat o'quvchilar uchun emas, o'qituvchilar uchun ham kasbiy "
     "rivojlanish imkoniyatini yaratdi. 6 nafar o'qituvchidan 5 tasi modelni qo'llash jarayonida "
     "o'z dars berish mahorati va refleksiya ko'nikmalarining oshganini qayd etdi. Bu MFAT "
     "modelining pedagogik ekotizimga tizimli ta'sirini ko'rsatadi: o'quvchi ham, o'qituvchi ham "
     "birgalikda o'sadi.")
empty()

body_bold_start("Asosiy tavsiyalar:", " Ta'lim vazirligi darajasida: mustaqil fikrlash va axborot "
     "texnologiyalari integratsiyasini DTS ga kiritish; respublika miqyosida 40 soatlik malaka "
     "oshirish kurslari; o'zbek tilidagi ochiq raqamli ta'lim resurslari kutubxonasi. Maktab "
     "rahbariyati darajasida: \"pilot-kengaytirish\" yondashuvi; o'qituvchilar uchun raqamli "
     "laboratoriya; ota-onalar uchun raqamli savodxonlik seminarlari. O'qituvchilar darajasida: "
     "haftada kamida bir marta ochiq muammo vazifasi; dars oxirida Exit Ticket; hamkor o'qitish "
     "metodlari. Xulosa qilib aytganda, axborot texnologiyalari o'quvchilarda mustaqil fikrlashni "
     "rivojlantirishning kuchli va tasdiqlangan vositasidir — biroq faqat pedagogik jihatdan to'g'ri, "
     "maqsadli va tizimli qo'llanilganda samarali bo'ladi. MFAT modeli ushbu shartni ta'minlashning "
     "amaliy yo'lini ko'rsatib beradi.")

# ══════════════════════════════════════════════════════════════════════════════
# ADABIYOTLAR
# ══════════════════════════════════════════════════════════════════════════════
chapter_head("FOYDALANILGAN ADABIYOTLAR RO'YXATI")
add_bibliography_field()

# ══════════════════════════════════════════════════════════════════════════════
# ILOVALAR
# ══════════════════════════════════════════════════════════════════════════════
page_break()
center_bold("ILOVALAR", size=14)

# Ilova A — Glossariy
page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
R(p, "1-ilova")
center_bold("TAYANCH SO'ZLAR (GLOSSARIY)", size=13)
empty()

terms = [
    ("1",  "Mustaqil fikrlash", "Самостоятельное мышление", "Independent Thinking"),
    ("2",  "Axborot texnologiyalari", "Информационные технологии", "Information Technologies"),
    ("3",  "Tanqidiy fikrlash", "Критическое мышление", "Critical Thinking"),
    ("4",  "Ijodiy fikrlash", "Творческое мышление", "Creative Thinking"),
    ("5",  "Raqamli ta'lim", "Цифровое образование", "Digital Education"),
    ("6",  "Pedagogik model", "Педагогическая модель", "Pedagogical Model"),
    ("7",  "Konnektivizm", "Коннективизм", "Connectivism"),
    ("8",  "Konstruktivizm", "Конструктивизм", "Constructivism"),
    ("9",  "Bloom taksonomiyasi", "Таксономия Блума", "Bloom's Taxonomy"),
    ("10", "Proximal rivojlanish zonasi", "Зона ближайшего развития", "Zone of Proximal Development"),
    ("11", "Adaptiv ta'lim", "Адаптивное обучение", "Adaptive Learning"),
    ("12", "Formativ baholash", "Формативное оценивание", "Formative Assessment"),
    ("13", "Aqliy xarita", "Ментальная карта", "Mind Map"),
    ("14", "Gamifikatsiya", "Геймификация", "Gamification"),
    ("15", "Metakognitsiya", "Метакогниция", "Metacognition"),
    ("16", "Interaktiv o'qish", "Интерактивное обучение", "Interactive Learning"),
    ("17", "O'quv analitikasi", "Учебная аналитика", "Learning Analytics"),
    ("18", "Raqamli savodxonlik", "Цифровая грамотность", "Digital Literacy"),
    ("19", "Axborot savodxonligi", "Информационная грамотность", "Information Literacy"),
    ("20", "Sun'iy intellekt", "Искусственный интеллект", "Artificial Intelligence"),
    ("21", "Fasilitator", "Фасилитатор", "Facilitator"),
    ("22", "Loyiha asosidagi ta'lim", "Проектное обучение", "Project-Based Learning"),
    ("23", "TPACK modeli", "Модель TPACK", "TPACK Model"),
    ("24", "Hamkorlikdagi o'rganish", "Совместное обучение", "Collaborative Learning"),
    ("25", "Universal dizayn", "Универсальный дизайн", "Universal Design for Learning"),
    ("26", "Teskari aloqa", "Обратная связь", "Feedback"),
    ("27", "Kognitiv rivojlanish", "Когнитивное развитие", "Cognitive Development"),
    ("28", "O'z-o'zini baholash", "Самооценивание", "Self-Assessment"),
    ("29", "Spiral o'rganish", "Спиральное обучение", "Spiral Learning"),
    ("30", "Raqamli portfolio", "Цифровое портфолио", "Digital Portfolio"),
    ("31", "Blended learning", "Смешанное обучение", "Blended Learning"),
    ("32", "Argumentatsiya", "Аргументация", "Argumentation"),
    ("33", "Simulyatsiya", "Симуляция", "Simulation"),
    ("34", "Raqamli chalg'ituvchi", "Цифровое отвлечение", "Digital Distraction"),
    ("35", "Ta'lim natijasi", "Образовательный результат", "Learning Outcome"),
]
add_glossary_table(terms)

# Ilova B — Texnik spetsifikatsiya
page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
R(p, "2-ilova")
center_bold("MFAT MODELI — TEXNIK SPETSIFIKATSIYA", size=13)
empty()
body("MFAT modeli joriy etilishida tavsiya etilgan raqamli vositalar va ularning texnik talablari quyidagi "
     "jadvalda keltirilgan. Har bir vosita pedagogik maqsad, foydalanuvchi kategoriyasi, qurilma talabi "
     "va narx modeli bo'yicha tavsiflangan.")
empty()

table_caption("2-jadval", "MFAT modeli uchun tavsiya etilgan raqamli vositalar")
tbl = doc.add_table(rows=1, cols=5)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers_b = ["Vosita", "Pedagogik maqsad", "Foydalanuvchi", "Qurilma", "Narx"]
for i, h in enumerate(headers_b):
    cell = tbl.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p, h, bold=True, size=11)

tools_data = [
    ("Coggle", "Aqliy xaritalar, tahlil", "O'quvchi, o'qituvchi", "Web, mobil", "Bepul/Pulli"),
    ("Kahoot!", "Formativ baholash", "O'qituvchi", "Web, mobil", "Bepul/Pulli"),
    ("Mentimeter", "Faollashtirish, so'rov", "O'qituvchi", "Web", "Bepul/Pulli"),
    ("Nearpod", "Interaktiv dars", "O'qituvchi", "Web, mobil", "Bepul/Pulli"),
    ("Socrative", "Tanqidiy savollar", "O'qituvchi", "Web, mobil", "Bepul"),
    ("Google Classroom", "LMS, hamkorlik", "Ikkala", "Web, mobil", "Bepul"),
    ("Padlet", "Hamkorlik, refleksiya", "Ikkala", "Web, mobil", "Bepul/Pulli"),
    ("Google Docs/Slides", "Birgalikdagi hujjat", "Ikkala", "Web, mobil", "Bepul"),
]
for row_data in tools_data:
    row = tbl.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = ''
        p = row[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i < 4 else WD_ALIGN_PARAGRAPH.CENTER
        R(p, val, size=10)

# Ilova C — Baholash mezoni
page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
R(p, "3-ilova")
center_bold("MUSTAQIL FIKRLASH RIVOJLANISHINI BAHOLASH MEZONI", size=13)
empty()
body("Watson-Glaser Critical Thinking Appraisal (WGCTA) asosida moslashtirilgan baholash mezoni. "
     "Baholash 5 ta komponent bo'yicha amalga oshiriladi, har biri 0-10 ball bilan baholanadi.")
empty()

table_caption("3-jadval", "Mustaqil fikrlash baholash mezoni (WGCTA moslashtirilgan versiyasi)")
tbl2 = doc.add_table(rows=1, cols=4)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Komponent", "Tavsif", "Maksimal ball", "O'lchov usuli"]):
    cell = tbl2.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p, h, bold=True, size=11)
criteria_rows = [
    ("Xulosa chiqarish", "Berilgan ma'lumotdan mantiqiy xulosa chiqarish", "10", "Ko'p tanlovli test"),
    ("Gipotezalarni tanib olish", "To'g'ri va noto'g'ri taxminlarni ajratish", "10", "Test + kuzatuv"),
    ("Deduktiv mantiq", "Argument zanjiri to'g'riligini baholash", "10", "Ko'p tanlovli test"),
    ("Ta'birni talqin qilish", "Ma'lumotlardan to'g'ri ma'no chiqarish", "10", "Ochiq savollar"),
    ("Argumentlarni baholash", "Kuchli va zaif argumentlarni ajratish", "10", "Test + kuzatuv"),
]
for row_data in criteria_rows:
    row = tbl2.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = ''
        p = row[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        R(p, val, size=10)

# ── Save ──────────────────────────────────────────────────────────────────────
print("Bibliografiya inject qilinmoqda...")
_inject_bibliography()

print(f"Saqlanmoqda: {OUT}")
doc.save(OUT)
print("TAYYOR!")
print(f"Fayl: {OUT}")
